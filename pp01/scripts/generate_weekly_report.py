from docx import Document
from docx.shared import Pt
import os

content = """
진행 내역 (6/8~6/12)

프로젝트 명: SAP ERP 도입 프로젝트
상태: 구현중 / 테스트 및 보완중

요약:
- 메인 화면과 연동되는 검사(ALV) 모듈의 편집 기능(입력/수정) 기본 구현 및 유효성 검사 로직 추가.
- OrderDetail 뷰의 데이터 바인딩 보강 및 추가 필드(EA 단위, 합계 수량) 표시 기능 구현.
- Value Help(자재/오더/출처) 프래그먼트 검색 및 선택 로직 개선 — 검색 결과 페이징 및 클리어 동작 안정화.
- 로컬 서비스 데이터 로드 및 에러 처리 강화(비동기 처리 체인 개선, 예외 메시지 처리 추가).
- 단위/통합 테스트 케이스 일부 작성 및 실행(특정 컨트롤러와 OPA 시나리오 확인).

상세 변경 및 작업 내용:
1) ALV Edit 기능 추가
- 검사 자재 데이터를 ALV Edit 형태로 표시하여 사용자가 수치와 자재를 직접 입력/수정 가능하도록 구현.
- 입력 형식 유효성 검사 추가: 비정상 값 입력 시 셀 색상 변경 및 저장(SAVE) 버튼 비활성화.
- 저장 시 동시성 문제 예방을 위해 저장 중 버튼 비활성화 및 실패 시 롤백/오류 알림 처리 구현.

2) OrderDetail 개선
- 상세화면에 EA 단위 필드 추가 및 기존 수량 계산 로직 보강.
- 추가 필드 바인딩 시 성능 저하를 줄이기 위해 데이터 바인딩 범위를 필요한 필드로 제한.

3) Value Help 및 검색 개선
- Material/Order/Source Value Help에서 키워드 검색 결과 정렬·페이징 추가.
- 선택 후 포커스 복귀 및 다중선택 시 상태 동기화 로직 보강.

4) 데이터 로드 및 오류 처리
- localService/sourceService의 데이터 호출 체인에서 Promise 처리와 에러 핸들링을 강화하여 네트워크 오류 시 재시도 또는 사용자 안내 제공.
- 메타데이터 변경에 따른 모델 매핑 업데이트(models.js 수정).

5) 테스트 및 문서화
- 주요 컨트롤러에 대해 유닛 테스트 추가(테스트 파일: test/unit/controller/Main.controller.js 기반).
- 통합 OPA 시나리오 일부 자동화로 네비게이션/화면 표시 흐름 검증.
- i18n에 신규 메시지 키 추가(검수 안내, 저장 성공/실패 메시지).

버그 수정
- ALV에서 특정 조작 후 발생하던 저장 불능 현상 수정(유효성 검사 플래그 초기화 문제 해결).
- 상세화면에서 일부 Null 참조로 인한 렌더링 오류 수정.

향후 계획 (우선순위)
- ALV Edit의 저장 유효성 검사 서버검증 연동 및 에러 세부 메시지 표시(다음주).
- 전체 워크플로우 동작 완료 후 회귀 테스트 및 마무리 검증.
- 기능 마무리 후 프로그램 단위 통합 테스트 후 배포 준비(프로그램명: ZC1MM0001).

예시(간단 설명)
- 검사 데이터 입력 시 허용 범위를 벗어나면 해당 셀을 붉게 표시하고 SAVE 기능을 못하게 막음.
- 입력 도중 BACK 등 다른 동작 발생 시 변경사항 유지/취소 동작 명확화.
"""

output_filename = os.path.join(os.getcwd(), 'weekly_report_2026-06-08_to_2026-06-12.docx')

doc = Document()

# Title
p = doc.add_paragraph()
run = p.add_run('주간 개발 일지 - 2026.06.08 ~ 2026.06.12')
run.bold = True
run.font.size = Pt(14)

# Add a blank line
doc.add_paragraph('')

# Add content paragraphs
for line in content.strip().split('\n'):
    doc.add_paragraph(line)

# Save
doc.save(output_filename)
print('Saved:', output_filename)
